# -*- coding: UTF-8 -*-
from docx import Document
import re
import os
import pandas as pd


path = raw_input(u"请输入文件目录:")
# path=unicode(path,'utf8')
dirs = os.listdir(path)
a1 = []
b1 = []
c1 = []
d1 = []
e1 = []
number=[]
ficc=[]
cost=[]
equi=[]
money=[]


def DelLastThreeChar(str):
    str_list=list(str)
    str_list.pop()
    str_list.pop()
    str_list.pop()
    return "".join(str_list)

# word 文档遍历
def dirlist(mainpath, allfilelist):
   filelist = os.listdir(mainpath)
   for filename in filelist:
      filepath = os.path.join(mainpath, filename)
      if os.path.isdir(filepath):
         dirlist(filepath, allfilelist)
      else:
         allfilelist.append(filepath)
   return allfilelist

dirs = dirlist(path, [])  # 获得目录下所有文件

allfile=[]
for dir in dirs:
    pat='结算通知书'
    m=re.search(pat,dir.decode('gbk').encode('utf8'))
    if m:
        allfile.append(dir) # 获得所有文档名中带有结算通知书的文件

#print allfile
print len(allfile)

findf=[]
for dir in allfile:
    try:
    #document=Document(dir.decode("utf8"))
        temp=0

    #document=Document(path+"\\"+dir)
        document = Document(dir)
        doc_new=Document()
        t=document.tables[0]

        for r in t.rows:
            for c in r.cells:
                doc_new.add_paragraph(c.text, style=None)
        a1.append(doc_new.paragraphs[1].text)
        findf.append(1)
        for inx, p in enumerate(doc_new.paragraphs):

            tt = p.text
            pat = '利率收益金额'.decode("utf8")
            m = re.search(pat, tt)
            if m:
                if temp == 0:
                    temp =1
                    b1.append(doc_new.paragraphs[inx + 1].text)
                    findf.append(2)

            pat = '权益收益金额'.decode("utf8")
            m = re.search(pat, tt)
            if m:
                c1.append(doc_new.paragraphs[inx + 1].text)
                findf.append(3)

            pat = '现金分红'.decode("utf8")
            m = re.search(pat, tt)
            if m:
                d1.append(doc_new.paragraphs[inx + 1].text)
                findf.append(4)

            pat = '交易费用'.decode("utf8")
            m = re.search(pat, tt)
            if m:
                e1.append(doc_new.paragraphs[inx + 1].text)
                findf.append(5)

    except Exception as e:
        print "Error: ",e #error 主要情况存在于windows word 自动保存文件

print len(a1)
print len(b1)
print len(c1)
print len(d1)
print len(e1)

#data_seesee=pd.DataFrame({"kan":findf})
#data_seesee.to_csv("D:\\error.csv",index= False)
head=["Number","Fixed Income","Equity","Cash","Trading"]
dataframe = pd.DataFrame({"Number": a1, "Fixed Income": b1, "Equity": c1, "Cash": d1, "Trading": e1})
# dataframe=pd.DataFrame(l,columns=head)
data1 = dataframe[head]
data1.to_csv(path+"\\result.csv", encoding='utf_8_sig', index=False)

# ''' ''' 部分代码原用于将excel中的中文去除，但由于部分文件会存在支付方重要信息中文，所以不执行。
'''
str1=data1['Number'][17].encode("utf-8")
print str1
print str1.split(": ")[1]
'''
'''
str = data1['Number'][0].encode("utf-8")
str1 = str.split()[2]
str2 = str1.split("：")[1]
str2=DelLastThreeChar(str2)
print str2
'''
'''
for index in data1.index:
    try:
        str=data1['Number'][index].encode("utf-8")
        str1=str.split()[2]
        str2=str1.split("：")[1]
    #str2=DelLastThreeChar(str2)
        number.append(str2)
    except Exception as e:
        print index,e
'''
'''
for index in data1.index:
    str=data1['Fixed Income'][index].encode("utf-8")
    str2=DelLastThreeChar(str)
    ficc.append(str2)

for index in data1.index:
    str=data1['Equity'][index].encode("utf-8")
    str2=DelLastThreeChar(str)
    equi.append(str2)

for index in data1.index:
    str=data1['Cash'][index].encode("utf-8")
    str2=DelLastThreeChar(str)
    money.append(str2)


for index in data1.index:
    str=data1['Trading'][index].encode("utf-8")
    str2=DelLastThreeChar(str)
    cost.append(str2)


dataframe_new = pd.DataFrame({"Number": a1, "Fixed Income": ficc, "Equity": equi, "Cash": money, "Trading": cost})
data_clean = dataframe_new[head]
print data_clean
'''
'''
str=data1['Number'][0].encode("utf-8")
str1=str.split( )[2]
print str1.split("：")[1]
'''
#data_clean.to_csv("D:\\222\\result.csv", encoding='utf_8_sig', index=False)


